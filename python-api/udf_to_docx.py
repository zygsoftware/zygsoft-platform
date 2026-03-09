import sys
import os
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.section import WD_ORIENT
import base64
import io
import zipfile

def is_zip_file(file_path):
    """Check if the file is a valid ZIP file"""
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            return True
    except zipfile.BadZipFile:
        return False

def get_alignment_style(alignment_value):
    """Convert alignment value from XML to Word alignment constant"""
    if alignment_value == "1":
        return WD_ALIGN_PARAGRAPH.CENTER
    elif alignment_value == "3":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment_value == "2":
        return WD_ALIGN_PARAGRAPH.RIGHT
    else:
        return WD_ALIGN_PARAGRAPH.LEFT

def convert_color(color_value):
    """Convert integer color value to RGBColor and return RGB values as a tuple"""
    if color_value is None:
        return None
    
    try:
        # Convert from negative integer to positive hex
        color_int = int(color_value)
        if color_int < 0:
            color_int = 0xFFFFFFFF + color_int + 1
        
        # Extract RGB values
        r = (color_int >> 16) & 0xFF
        g = (color_int >> 8) & 0xFF
        b = color_int & 0xFF
        
        # Return both the RGBColor object and the RGB values
        return (RGBColor(r, g, b), (r, g, b))
    except (ValueError, TypeError):
        return None

def add_page_number(paragraph):
    """Add a page number field to the paragraph"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def set_cell_background(cell, color_info):
    """Set cell background color using direct XML manipulation"""
    color_obj, rgb_values = color_info
    shading_elm = OxmlElement('w:shd')
    # Convert RGB to hex
    hex_color = f"{rgb_values[0]:02X}{rgb_values[1]:02X}{rgb_values[2]:02X}"
    shading_elm.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def process_background_image(document, bg_image_data, bg_image_source, output_file):
    """Process background image data and add to document background"""
    if bg_image_data:
        try:
            # Decode base64 image data
            image_bytes = base64.b64decode(bg_image_data)
            
            # Save the image to a file next to the output document
            temp_img_path = os.path.join(os.path.dirname(output_file), os.path.splitext(os.path.basename(output_file))[0] + "_background.png")
            with open(temp_img_path, "wb") as img_file:
                img_file.write(image_bytes)
            
            # Unfortunately python-docx doesn't directly support setting background images
            # This requires OOXML manipulation, which is beyond the scope of this script
            print(f"Background image saved to {temp_img_path}. Please manually set it as document background in Word.")
            return True
        except Exception as e:
            print(f"Error processing background image data: {e}")
    elif bg_image_source:
        print(f"Background image source path: {bg_image_source}. Please manually set it as document background in Word.")
    return False

def udf_to_docx(udf_file, docx_file):
    root = None
    
    # Check if the file is a ZIP file
    if is_zip_file(udf_file):
        # Process as a ZIP file
        with zipfile.ZipFile(udf_file, 'r') as z:
            if 'content.xml' in z.namelist():
                with z.open('content.xml') as content_file:
                    tree = ET.parse(content_file, parser=ET.XMLParser(encoding='utf-8'))
                    root = tree.getroot()
            else:
                print("The 'content.xml' file could not be found in the UDF file.")
                exit()
    else:
        # Process as an XML file directly
        try:
            tree = ET.parse(udf_file, parser=ET.XMLParser(encoding='utf-8'))
            root = tree.getroot()
        except ET.ParseError:
            print(f"The file {udf_file} is neither a valid ZIP nor a valid XML file.")
            exit()

    if root is None:
        print("Failed to parse the file.")
        exit()

    # Create a new Word document
    document = Document()
    
    # Ensure default headers and footers are created for all sections
    for section in document.sections:
        section.different_first_page = False
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    # Retrieve content text
    content_element = root.find('content')
    if content_element is not None:
        content_text = content_element.text
        if content_text.startswith('<![CDATA[') and content_text.endswith(']]>'):
            content_text = content_text[9:-3]
    else:
        print("'content' could not be found in the XML.")
        exit()

    # Extract page properties
    properties_element = root.find('properties')
    page_format = properties_element.find('pageFormat') if properties_element is not None else None
    
    # Get page margins
    if page_format is not None:
        left_margin = float(page_format.get('leftMargin', '42.5')) / 72 * Inches(1).pt  # Convert to Word points
        right_margin = float(page_format.get('rightMargin', '42.5')) / 72 * Inches(1).pt
        top_margin = float(page_format.get('topMargin', '42.5')) / 72 * Inches(1).pt
        bottom_margin = float(page_format.get('bottomMargin', '42.5')) / 72 * Inches(1).pt
        
        # Set page margins for each section
        for section in document.sections:
            section.left_margin = Pt(left_margin)
            section.right_margin = Pt(right_margin)
            section.top_margin = Pt(top_margin)
            section.bottom_margin = Pt(bottom_margin)
            
            # Set page orientation
            orientation = page_format.get('paperOrientation', '1')
            if orientation == '2':  # Landscape
                section.orientation = WD_ORIENT.LANDSCAPE
            else:  # Portrait
                section.orientation = WD_ORIENT.PORTRAIT
    
    # Get background image if available
    if properties_element is not None:
        bg_image_elem = properties_element.find('bgImage')
        if bg_image_elem is not None:
            bg_image_data = bg_image_elem.get('bgImageData')
            bg_image_source = bg_image_elem.get('bgImageSource')
            process_background_image(document, bg_image_data, bg_image_source, docx_file)

    # Process the 'elements' section
    elements_element = root.find('elements')
    if elements_element is not None:
        # Get header and footer elements
        header_element = elements_element.find('header')
        footer_element = elements_element.find('footer')
        
        # Process header
        if header_element is not None:
            # Get the header from the first section
            section = document.sections[0]
            header = section.header
            
            # Clear existing header paragraphs
            for p in header.paragraphs:
                p._element.getparent().remove(p._element)
                p._p = None
                p._element = None
            
            # Create new header paragraph
            header_para = header.add_paragraph()
            
            # Set header background color
            header_color_result = convert_color(header_element.get('background'))
            if header_color_result:
                header_bg_color, rgb_values = header_color_result
                # Save background color info for manual formatting
                print(f"Header background color: RGB({rgb_values[0]}, {rgb_values[1]}, {rgb_values[2]}) - Please set it manually in Word.")
            
            # Process header paragraphs
            for para_elem in header_element.findall('paragraph'):
                if para_elem is not header_element.findall('paragraph')[0]:
                    header_para = header.add_paragraph()
                
                # Set alignment
                alignment = para_elem.get('Alignment', '0')
                header_para.alignment = get_alignment_style(alignment)
                
                # Process content
                for child in para_elem:
                    if child.tag == 'content':
                        start_offset = int(child.get('startOffset', '0'))
                        length = int(child.get('length', '0'))
                        text = content_text[start_offset:start_offset+length]
                        
                        run = header_para.add_run(text)
                        
                        # Always use DejaVuSerif as default font
                        run.font.name = "DejaVuSerif"
                        
                        # Set font size
                        size = child.get('size')
                        if size:
                            run.font.size = Pt(float(size))
                        
                        # Set formatting
                        bold = child.get('bold', 'false') == 'true'
                        italic = child.get('italic', 'false') == 'true'
                        underline = child.get('underline', 'false') == 'true'
                        
                        run.bold = bold
                        run.italic = italic
                        if underline:
                            run.underline = WD_UNDERLINE.SINGLE
                        
                        # Set color
                        foreground_result = convert_color(child.get('foreground'))
                        if foreground_result:
                            # Extract just the RGBColor object, not the tuple
                            foreground = foreground_result[0]
                            run.font.color.rgb = foreground
        
        # Process footer
        if footer_element is not None:
            # Get the footer from the first section
            section = document.sections[0]
            footer = section.footer
            
            # Clear existing footer paragraphs
            for p in footer.paragraphs:
                p._element.getparent().remove(p._element)
                p._p = None
                p._element = None
            
            # Create new footer paragraph
            footer_para = footer.add_paragraph()
            
            # Process footer background color
            footer_color_result = convert_color(footer_element.get('background'))
            if footer_color_result:
                footer_bg_color, rgb_values = footer_color_result
                print(f"Footer background color: RGB({rgb_values[0]}, {rgb_values[1]}, {rgb_values[2]}) - Please set it manually in Word.")
            
            # Process footer paragraphs
            for para_elem in footer_element.findall('paragraph'):
                if para_elem is not footer_element.findall('paragraph')[0]:
                    footer_para = footer.add_paragraph()
                
                # Set alignment
                alignment = para_elem.get('Alignment', '0')
                footer_para.alignment = get_alignment_style(alignment)
                
                # Process content
                for child in para_elem:
                    if child.tag == 'content':
                        start_offset = int(child.get('startOffset', '0'))
                        length = int(child.get('length', '0'))
                        text = content_text[start_offset:start_offset+length]
                        
                        run = footer_para.add_run(text)
                        
                        # Always use DejaVuSerif as default font
                        run.font.name = "DejaVuSerif"
                        
                        # Set font size
                        size = child.get('size')
                        if size:
                            run.font.size = Pt(float(size))
                        
                        # Set formatting
                        bold = child.get('bold', 'false') == 'true'
                        italic = child.get('italic', 'false') == 'true'
                        underline = child.get('underline', 'false') == 'true'
                        
                        run.bold = bold
                        run.italic = italic
                        if underline:
                            run.underline = WD_UNDERLINE.SINGLE
                        
                        # Set color
                        foreground_result = convert_color(child.get('foreground'))
                        if foreground_result:
                            # Extract just the RGBColor object, not the tuple
                            foreground = foreground_result[0]
                            run.font.color.rgb = foreground
                            
            # Add page number if needed (optional)
            # This can be uncommented if page numbers are required in the footer
            # add_page_number(footer_para)
        
        # Process each element in the document body
        for elem in elements_element:
            if elem.tag == 'paragraph':
                # Create the paragraph
                paragraph = document.add_paragraph()

                # Set paragraph alignment
                alignment = elem.get('Alignment', '0')
                paragraph.alignment = get_alignment_style(alignment)
                
                # Set paragraph indentation
                left_indent = elem.get('LeftIndent')
                right_indent = elem.get('RightIndent')
                first_line_indent = elem.get('FirstLineIndent')
                
                if left_indent:
                    paragraph.paragraph_format.left_indent = Pt(float(left_indent))
                if right_indent:
                    paragraph.paragraph_format.right_indent = Pt(float(right_indent))
                if first_line_indent:
                    paragraph.paragraph_format.first_line_indent = Pt(float(first_line_indent))
                
                # Set line spacing
                line_spacing = elem.get('LineSpacing')
                if line_spacing:
                    paragraph.paragraph_format.line_spacing = float(line_spacing)
                
                # Set list properties
                bulleted = elem.get('Bulleted', 'false') == 'true'
                numbered = elem.get('Numbered', 'false') == 'true'
                
                if bulleted:
                    # Set bullet list style
                    paragraph.style = 'List Bullet'
                elif numbered:
                    # Set numbered list style
                    paragraph.style = 'List Number'

                # Get paragraph-level font properties (used as defaults)
                para_size = elem.get('size')
                para_family = elem.get('family', 'DejaVuSerif')  # Default font
                para_default_size = float(para_size) if para_size else 12.0

                # Process the paragraph content
                for child in elem:
                    if child.tag == 'content':
                        # Get and format the text
                        start_offset = int(child.get('startOffset', '0'))
                        length = int(child.get('length', '0'))
                        text = content_text[start_offset:start_offset+length]

                        run = paragraph.add_run(text)

                        # Set font - use content's font if specified, otherwise paragraph default
                        content_family = child.get('family')
                        if content_family:
                            run.font.name = content_family
                        else:
                            run.font.name = para_family

                        # Set the font size - use content's size if specified, otherwise paragraph default
                        size = child.get('size')
                        if size:
                            try:
                                run.font.size = Pt(float(size))
                            except (ValueError, TypeError):
                                run.font.size = Pt(para_default_size)
                        else:
                            run.font.size = Pt(para_default_size)

                        # Set text formatting
                        bold = child.get('bold', 'false') == 'true'
                        italic = child.get('italic', 'false') == 'true'
                        underline = child.get('underline', 'false') == 'true'
                        
                        run.bold = bold
                        run.italic = italic
                        if underline:
                            run.underline = WD_UNDERLINE.SINGLE

                        # Set the color
                        foreground_result = convert_color(child.get('foreground'))
                        if foreground_result:
                            # Extract just the RGBColor object, not the tuple
                            foreground = foreground_result[0]
                            run.font.color.rgb = foreground

                    elif child.tag == 'field':
                        # Process field element (labels like DAVACI, VEKİLİ, etc.)
                        field_name = child.get('fieldName', '')
                        
                        # Get the text from the content buffer if startOffset and length are provided
                        if child.get('startOffset') and child.get('length'):
                            start_offset = int(child.get('startOffset', '0'))
                            length = int(child.get('length', '0'))
                            field_text = content_text[start_offset:start_offset+length]
                        else:
                            # Use the fieldName as fallback
                            field_text = field_name
                        
                        run = paragraph.add_run(field_text)
                        
                        # Always use DejaVuSerif as font
                        run.font.name = "DejaVuSerif"
                        
                        # Set formatting based on attributes
                        run.bold = child.get('bold', 'false') == 'true'
                        run.italic = child.get('italic', 'false') == 'true'
                        if child.get('underline', 'false') == 'true':
                            run.underline = WD_UNDERLINE.SINGLE
                        
                        # Set color if available
                        foreground_result = convert_color(child.get('foreground'))
                        if foreground_result:
                            # Extract just the RGBColor object, not the tuple
                            foreground = foreground_result[0]
                            run.font.color.rgb = foreground
                            
                    elif child.tag == 'space':
                        # Add a space
                        run = paragraph.add_run(" ")
                        
                    elif child.tag == 'image':
                        # Add an image
                        image_data = child.get('imageData')
                        if image_data:
                            image_bytes = base64.b64decode(image_data)
                            image_stream = io.BytesIO(image_bytes)
                            run = paragraph.add_run()
                            run.add_picture(image_stream)
                            
            elif elem.tag == 'page-break':
                # Add page break
                document.add_page_break()
                
            elif elem.tag == 'table':
                # Create the table
                column_count = int(elem.get('columnCount', '1'))
                rows = elem.findall('row')
                
                # Get column widths if specified
                col_widths = []
                col_spans = elem.get('columnSpans', '')
                if col_spans:
                    try:
                        col_spans_list = col_spans.split(',')
                        if len(col_spans_list) == column_count:
                            for span in col_spans_list:
                                col_widths.append(Pt(float(span)))
                    except (ValueError, IndexError):
                        col_widths = []
                
                # Create table
                table = document.add_table(rows=len(rows), cols=column_count)
                
                # Set border style
                border_style = elem.get('border', 'borderCell')
                if border_style in ['borderCell', 'border']:
                    # Add borders to all cells
                    table.style = 'Table Grid'
                elif border_style == 'borderOuter':
                    # Only outer borders
                    table.style = 'Table Grid'
                    # Would need more complex XML manipulation to properly implement 'borderOuter'
                    
                # Process table rows and cells
                for row_idx, row in enumerate(rows):
                    # Set row height if specified
                    row_height = row.get('height_min')
                    if row_height:
                        table.rows[row_idx].height = Pt(float(row_height) * 72)  # Convert to points
                        
                    cells = row.findall('cell')
                    for col_idx, cell in enumerate(cells):
                        # Ensure we don't exceed column count
                        if col_idx >= column_count:
                            continue
                            
                        # Get the table cell
                        table_cell = table.rows[row_idx].cells[col_idx]
                        
                        # Process cell paragraphs
                        paragraphs = cell.findall('paragraph')
                        
                        # Use existing paragraph if possible
                        cell_paragraph = table_cell.paragraphs[0] if table_cell.paragraphs else table_cell.add_paragraph()
                        
                        for para_idx, para in enumerate(paragraphs):
                            # Add a new paragraph for subsequent paragraphs
                            if para_idx > 0:
                                cell_paragraph = table_cell.add_paragraph()
                                
                            # Set paragraph alignment
                            alignment = para.get('Alignment', '0')
                            cell_paragraph.alignment = get_alignment_style(alignment)
                            
                            # Set paragraph indentation
                            left_indent = para.get('LeftIndent')
                            right_indent = para.get('RightIndent')
                            
                            if left_indent:
                                cell_paragraph.paragraph_format.left_indent = Pt(float(left_indent))
                            if right_indent:
                                cell_paragraph.paragraph_format.right_indent = Pt(float(right_indent))
                            
                            # Process paragraph content
                            for child in para:
                                if child.tag == 'content':
                                    # Get and format the text
                                    start_offset = int(child.get('startOffset', '0'))
                                    length = int(child.get('length', '0'))
                                    text = content_text[start_offset:start_offset+length]

                                    run = cell_paragraph.add_run(text)

                                    # Always use DejaVuSerif as font
                                    run.font.name = "DejaVuSerif"

                                    # Set the font size
                                    size = child.get('size')
                                    if size:
                                        run.font.size = Pt(float(size))

                                    # Set text formatting
                                    bold = child.get('bold', 'false') == 'true'
                                    italic = child.get('italic', 'false') == 'true'
                                    underline = child.get('underline', 'false') == 'true'
                                    
                                    run.bold = bold
                                    run.italic = italic
                                    if underline:
                                        run.underline = WD_UNDERLINE.SINGLE

                                    # Set the color
                                    foreground_result = convert_color(child.get('foreground'))
                                    if foreground_result:
                                        # Extract just the RGBColor object, not the tuple
                                        foreground = foreground_result[0]
                                        run.font.color.rgb = foreground

                                elif child.tag == 'field':
                                    # Process field element
                                    field_name = child.get('fieldName', '')
                                    
                                    # Get the text from the content buffer
                                    if child.get('startOffset') and child.get('length'):
                                        start_offset = int(child.get('startOffset', '0'))
                                        length = int(child.get('length', '0'))
                                        field_text = content_text[start_offset:start_offset+length]
                                    else:
                                        # Use the fieldName as fallback
                                        field_text = field_name
                                    
                                    run = cell_paragraph.add_run(field_text)
                                    
                                    # Always use DejaVuSerif as font
                                    run.font.name = "DejaVuSerif"
                                    
                                    # Set formatting
                                    run.bold = child.get('bold', 'false') == 'true'
                                    run.italic = child.get('italic', 'false') == 'true'
                                    if child.get('underline', 'false') == 'true':
                                        run.underline = WD_UNDERLINE.SINGLE
                                    
                                    # Set color if available
                                    foreground_result = convert_color(child.get('foreground'))
                                    if foreground_result:
                                        # Extract just the RGBColor object, not the tuple
                                        foreground = foreground_result[0]
                                        run.font.color.rgb = foreground
                                        
                                elif child.tag == 'space':
                                    # Add a space
                                    cell_paragraph.add_run(" ")
                                elif child.tag == 'image':
                                    # Add an image
                                    image_data = child.get('imageData')
                                    if image_data:
                                        try:
                                            image_bytes = base64.b64decode(image_data)
                                            image_stream = io.BytesIO(image_bytes)
                                            run = cell_paragraph.add_run()
                                            run.add_picture(image_stream)
                                        except Exception as e:
                                            print(f"Error processing image in table: {e}")
                                            cell_paragraph.add_run("[GÖRSEL]")
    else:
        print("'elements' could not be found in the XML.")
        exit()

    # Save the document
    document.save(docx_file)
    print(f"DOCX file created: {docx_file}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python udf_to_docx.py input.udf")
        exit()

    udf_file = sys.argv[1]

    if not os.path.isfile(udf_file):
        print(f"Input file not found: {udf_file}")
        exit()

    filename, ext = os.path.splitext(udf_file)

    if ext.lower() == '.udf':
        docx_file = filename + '.docx'
        udf_to_docx(udf_file, docx_file)
    else:
        print("Please provide a .udf file.")

if __name__ == '__main__':
    main()