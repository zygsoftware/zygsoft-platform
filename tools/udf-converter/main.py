import re
import zipfile
from docx import Document
from paragraph_processor import process_paragraph, is_pseudo_list_paragraph, get_list_kind, is_numbered_heading_paragraph
from table_processor import process_table

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _cdata_safe(s: str) -> str:
    return (s or "").replace("]]>", "]]]]><![CDATA[>")

def _normalize_newlines(s: str) -> str:
    return (s or "").replace("\r\n", "\n").replace("\r", "\n")

def _is_real_list_paragraph(paragraph) -> bool:
    try:
        if paragraph is None:
            return False
        numPr = paragraph._element.find('.//{%s}numPr' % W_NS)
        return numPr is not None
    except Exception:
        return False

def _is_pseudo_list_paragraph(paragraph) -> bool:
    try:
        if paragraph is None:
            return False
        return is_pseudo_list_paragraph(paragraph)
    except Exception:
        return False

def _is_list_like_paragraph(paragraph) -> bool:
    return _is_real_list_paragraph(paragraph) or _is_pseudo_list_paragraph(paragraph)

def _is_bullet_list_paragraph(paragraph, document=None) -> bool:
    try:
        return get_list_kind(paragraph, document) == "bullet"
    except Exception:
        return False

def _is_heading_paragraph(paragraph) -> bool:
    try:
        return is_numbered_heading_paragraph(paragraph)
    except Exception:
        return False

def _is_empty_paragraph(paragraph) -> bool:
    """
    Görünmeyen boşluklar/U+00A0 vs. için sağlam kontrol.
    """
    try:
        if paragraph is None:
            return True
        t = paragraph.text or ""
        t = t.replace("\u00a0", " ").replace("\u200b", "").strip()
        return not t
    except Exception:
        return True

def _emit_blank_paragraph(current_offset: int):
    """
    Tam bos paragraf yerine tek NBSP kullan.
    Bu, hem UDF tarafinda bos satiri korur hem PDF render'da fazladan satir kirilmasi yaratmaz.
    """
    placeholder = "\u00a0"
    para_text = placeholder

    para_element = (
        f'<paragraph Alignment="0">'
        f'<content startOffset="{current_offset}" length="{len(placeholder)}" '
        f'family="Times New Roman" size="11" />'
        f'</paragraph>'
    )

    return para_text, para_element, current_offset + len(placeholder)

def _find_next_meaningful_kind(body_elems, kind_flags, p_map, start_idx):
    """
    start_idx'ten sonra boş paragraf(ları) atlayıp ilk anlamlı elemana bakar.
    Dönen değer: True/False (o anlamlı eleman liste mi?)
    """
    for j in range(start_idx + 1, len(body_elems)):
        el = body_elems[j]
        if el.tag.endswith('p'):
            p = p_map.get(el)
            if p is None:
                continue
            if _is_empty_paragraph(p):
                continue
            return kind_flags[j]
        else:
            return None
    return None

def _find_prev_meaningful_kind(body_elems, kind_flags, p_map, start_idx):
    """
    start_idx'ten önce boş paragraf(ları) atlayıp ilk anlamlı elemana bakar.
    """
    for j in range(start_idx - 1, -1, -1):
        el = body_elems[j]
        if el.tag.endswith('p'):
            p = p_map.get(el)
            if p is None:
                continue
            if _is_empty_paragraph(p):
                continue
            return kind_flags[j]
        else:
            return None
    return None

def _has_next_meaningful_element(body_elems, p_map, start_idx):
    for j in range(start_idx + 1, len(body_elems)):
        el = body_elems[j]
        if el.tag.endswith('p'):
            p = p_map.get(el)
            if p is None:
                continue
            if _is_empty_paragraph(p):
                continue
            return True
        return True
    return False

def _process_document_elements(docx_path: str):
    document = Document(docx_path)
    process_paragraph.list_counters = {}

    elements = []
    content = []
    current_offset = 0

    p_map = {p._element: p for p in document.paragraphs}
    t_map = {t._element: t for t in document.tables}

    body_elems = list(document.element.body)

    # paragraph kind: bullet / number / None
    kind_flags = []
    for el in body_elems:
        if el.tag.endswith('p'):
            p = p_map.get(el)
            kind_flags.append(get_list_kind(p, document) if p else None)
        else:
            kind_flags.append(None)

    def _append_blank():
        nonlocal current_offset
        blank_text, blank_el, new_off = _emit_blank_paragraph(current_offset)
        elements.append(blank_el)
        content.append(blank_text)
        current_offset = new_off

    for idx, el in enumerate(body_elems):
        kind_now = kind_flags[idx]
        prev_kind = _find_prev_meaningful_kind(body_elems, kind_flags, p_map, idx)
        next_kind = _find_next_meaningful_kind(body_elems, kind_flags, p_map, idx)
        is_bullet_now = kind_now == "bullet"
        is_bullet_prev_meaningful = prev_kind == "bullet"
        is_bullet_next_meaningful = next_kind == "bullet"
        should_add_blank_before_next_bullet = False

        if el.tag.endswith('p'):
            paragraph = p_map.get(el)
            if paragraph:
                is_heading_now = _is_heading_paragraph(paragraph)
                if _is_empty_paragraph(paragraph) and (is_bullet_prev_meaningful or is_bullet_next_meaningful):
                    continue

                suppress_space_after = is_bullet_now or is_bullet_next_meaningful or is_heading_now
                para_text, para_element, new_offset = process_paragraph(
                    paragraph,
                    document,
                    current_offset,
                    {"suppress_space_after": suppress_space_after},
                )
                elements.append(para_element)
                content.append(para_text)
                current_offset = new_offset
                should_add_blank_before_next_bullet = (
                    not is_bullet_now and
                    is_bullet_next_meaningful and
                    not _is_empty_paragraph(paragraph)
                )

        elif el.tag.endswith('tbl'):
            table = t_map.get(el)
            if table:
                table_text, table_element, new_offset = process_table(table, document, current_offset)
                elements.append(table_element)
                content.append(table_text)
                current_offset = new_offset
                should_add_blank_before_next_bullet = is_bullet_next_meaningful

        if should_add_blank_before_next_bullet:
            _append_blank()

        # Bullet liste bloğu bitişi: şimdi bullet, sonraki anlamlı eleman bullet değil
        if is_bullet_now and not is_bullet_next_meaningful:
            _append_blank()
            continue

        if el.tag.endswith('p'):
            paragraph = p_map.get(el)
            if paragraph and _is_heading_paragraph(paragraph) and _has_next_meaningful_element(body_elems, p_map, idx) and not is_bullet_next_meaningful:
                _append_blank()

    return "\n".join(elements), "".join(content)

def convert_docx_to_udf(docx_file: str, udf_file: str, template_xml_path: str = "calisanudfcontent.xml"):
    import logging
    _log = logging.getLogger(__name__)
    _log.info("[main.convert_docx_to_udf] template_xml_path=%s", template_xml_path)

    elements_xml, content_text = _process_document_elements(docx_file)

    with open(template_xml_path, "r", encoding="utf-8", errors="ignore") as f:
        tpl_xml = f.read()

    tpl_bg_match = re.search(r'bgImageData\s*=\s*["\']([^"\']*)["\']', tpl_xml, re.I)
    tpl_bg_len = len(tpl_bg_match.group(1)) if tpl_bg_match else 0
    _log.info("[main.convert_docx_to_udf] template bgImageData length=%d", tpl_bg_len)

    xml = tpl_xml

    # --- (A) content ---
    if re.search(r'(<content><!\[CDATA\[)(.*?)(\]\]></content>)', xml, flags=re.S | re.I):
        xml = re.sub(
            r'(<content><!\[CDATA\[)(.*?)(\]\]></content>)',
            lambda m: f"{m.group(1)}{_cdata_safe(content_text)}{m.group(3)}",
            xml, flags=re.S | re.I
        )
    else:
        xml = re.sub(
            r'(<properties>)',
            f'<content><![CDATA[{_cdata_safe(content_text)}]]></content>\n\\1',
            xml, count=1, flags=re.I
        )

    xml = re.sub(
        r'(</content>)(<properties>)',
        r'\1\n\2',
        xml, count=1, flags=re.I
    )

    # --- (B) elements ---
    if re.search(r'(<elements\b[^>]*>)(.*?)(</elements>)', xml, flags=re.S | re.I):
        xml = re.sub(
            r'(<elements\b[^>]*>)(.*?)(</elements>)',
            lambda m: f"{m.group(1)}{elements_xml}{m.group(3)}",
            xml, flags=re.S | re.I
        )
    else:
        xml = re.sub(
            r'(</content>)',
            f'\\1\n<elements resolver="hvl-default">\n{elements_xml}\n</elements>',
            xml, count=1, flags=re.I
        )

    out_bg_match = re.search(r'bgImageData\s*=\s*["\']([^"\']*)["\']', xml, re.I)
    out_bg_len = len(out_bg_match.group(1)) if out_bg_match else 0
    _log.info("[main.convert_docx_to_udf] final xml bgImageData length=%d before write", out_bg_len)
    if tpl_bg_len > 0 and out_bg_len == 0:
        _log.warning("[main.convert_docx_to_udf] bgImageData LOST: template had %d chars, output has 0", tpl_bg_len)

    with zipfile.ZipFile(udf_file, "w", compression=zipfile.ZIP_STORED) as z:
        z.writestr("content.xml", xml.encode("utf-8"))

def main(docx_file: str, udf_file: str):
    convert_docx_to_udf(docx_file, udf_file)
    print(f"Conversion completed: {docx_file} -> {udf_file}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python main.py input.docx [output.udf]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '.udf')
    main(input_file, output_file)
